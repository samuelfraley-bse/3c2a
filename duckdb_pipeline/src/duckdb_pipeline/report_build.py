"""Chart rendering + docx assembly for the weekly coach report.

Ports the styling/table helpers from `analysis/plot_charts.py` (font
registration, Foothill palette, axis styling) and
`analysis/build_preview_docx.py` (table-building helpers, `ordinal()`).
Charts are matplotlib PNGs embedded via python-docx `doc.add_picture(...)`,
matching the established `reports/build_report.py` precedent -- no reason to
introduce a different charting approach for a coach-facing deliverable.

Presentation only: everything here consumes the plain dicts/lists returned
by `report_data.py` and produces a `.docx` file. No SQL lives in this module.
"""

from __future__ import annotations

import os
from pathlib import Path

from . import report_data as rd

HEAVY = "United Serif Rg Hv"
MEDIUM = "United Serif Rg Md"

RED = "#A61E2F"
BLACK = "#000000"
GRAY = "#D0D1CE"

_FONT_PATHS = {
    "heavy": r"C:\Users\sffra\Downloads\united-serif-font\United Serif Reg Heavy\United Serif Reg Heavy.otf",
    "medium": r"C:\Users\sffra\Downloads\united-serif-reg-medium_6KssY\United Serif Reg Medium\United Serif Reg Medium.otf",
}


def _load_chart_fonts():
    """Registers the house fonts if present on this machine; falls back to
    matplotlib defaults otherwise (same guarded pattern as analysis/plot_charts.py).
    """
    import matplotlib.font_manager as fm

    for path in _FONT_PATHS.values():
        if os.path.exists(path):
            fm.fontManager.addfont(path)
    font_heavy = fm.FontProperties(fname=_FONT_PATHS["heavy"])
    font_medium = fm.FontProperties(fname=_FONT_PATHS["medium"])
    return font_heavy, font_medium


def _style_ax(ax, font_heavy, font_name, title, xlabel="", ylabel=""):
    ax.set_title(title, fontproperties=font_heavy, fontsize=14, color=BLACK, pad=10)
    ax.set_xlabel(xlabel, fontfamily=font_name, fontsize=10, color=BLACK)
    ax.set_ylabel(ylabel, fontfamily=font_name, fontsize=10, color=BLACK)
    ax.tick_params(colors=BLACK, labelsize=9)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontfamily(font_name)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(GRAY)
    ax.spines["bottom"].set_color(GRAY)
    ax.set_facecolor("white")
    ax.yaxis.grid(True, color=GRAY, linewidth=0.5, linestyle="--")
    ax.set_axisbelow(True)


def render_success_rate_by_down_chart(
    down_rows: list[dict[str, object]],
    team: str,
    opponent: str,
    out_path: str,
) -> str:
    """Success Rate by Down (1st-4th): one line for the featured team's
    offense, one for the opponent's defense. Down is ordinal, so a line
    chart reads the trend across downs more directly than a bar-per-down
    table would.
    """
    import matplotlib.pyplot as plt

    font_heavy, font_medium = _load_chart_fonts()
    font_name = font_medium.get_name()

    downs = [row["down"] for row in down_rows]
    off_rates = [row["team_off_success_rate"] for row in down_rows]
    def_rates = [row["opp_def_success_rate"] for row in down_rows]

    fig, ax = plt.subplots(figsize=(6.2, 3.6))
    ax.plot(downs, off_rates, color=RED, marker="o", linewidth=2, label=f"{team} Offense")
    ax.plot(downs, def_rates, color=BLACK, marker="s", linewidth=2, label=f"{opponent} Defense")
    ax.set_xticks(downs)
    ax.set_xticklabels([f"{d}{'st' if d == 1 else 'nd' if d == 2 else 'rd' if d == 3 else 'th'}" for d in downs])
    ax.legend(prop=font_medium, frameon=False)
    _style_ax(ax, font_heavy, font_name, "Success Rate by Down", ylabel="Success Rate %")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


# (trend_off_column, trend_def_column, subplot_title, y_label)
# Swapping which metrics headline the weekly-trends page is a one-line edit
# to this list -- no changes needed to the page-assembly code below.
DEFAULT_TREND_METRICS = [
    ("success_rate_off", "success_rate_def", "Success Rate", "Success Rate %"),
    ("explosive_rate_off", "explosive_rate_def", "Explosive Play Rate", "Explosive %"),
    ("third_down_conv_off", "third_down_conv_def", "3rd Down Conversion Rate", "Conversion %"),
]


def render_weekly_trends_chart(
    trend_rows: list[dict[str, object]],
    team: str,
    out_path: str,
    metrics: list[tuple[str, str, str, str]] = DEFAULT_TREND_METRICS,
) -> str:
    """One figure, one subplot per configured metric (default: 3), each with
    an offense line and a defense-faced line across the season's games so
    far. Stays a single embedded image regardless of how many metrics are
    configured, to respect the report's chart-count budget.
    """
    import matplotlib.pyplot as plt

    font_heavy, font_medium = _load_chart_fonts()
    font_name = font_medium.get_name()

    opponents = [row["opponent"] for row in trend_rows]
    fig, axes = plt.subplots(1, len(metrics), figsize=(4.2 * len(metrics), 3.6))
    if len(metrics) == 1:
        axes = [axes]

    for ax, (off_col, def_col, title, ylabel) in zip(axes, metrics):
        off_vals = [row.get(off_col) for row in trend_rows]
        def_vals = [row.get(def_col) for row in trend_rows]
        ax.plot(opponents, off_vals, color=RED, marker="o", linewidth=2, label="Offense")
        ax.plot(opponents, def_vals, color=BLACK, marker="s", linewidth=2, label="Defense Faced")
        ax.tick_params(axis="x", labelrotation=90)
        _style_ax(ax, font_heavy, font_name, title, ylabel=ylabel)

    axes[0].legend(prop=font_medium, frameon=False, loc="best")
    fig.suptitle(f"{team} Weekly Trends", fontproperties=font_heavy, fontsize=15, y=1.03)
    fig.tight_layout()

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def ordinal(n: int) -> str:
    n = int(n)
    suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10 if n % 100 not in (11, 12, 13) else 0, "th")
    return f"{n}{suffix}"


def fmt(value: object, kind: str) -> str:
    if value is None:
        return "--"
    if kind == "int":
        return str(int(round(float(value))))
    if kind == "pct":
        return f"{float(value):.1f}%"
    if kind == "f1":
        return f"{float(value):.1f}"
    if kind == "f2":
        return f"{float(value):.2f}"
    return str(value)


def _fmt_cell(value: object, rank: object, kind: str) -> str:
    text = fmt(value, kind)
    if rank is None:
        return text
    return f"{text} ({ordinal(rank)})"


def _set_cell(cell, text, font_name=HEAVY, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.bold = bold


def _add_subheading(doc, text, size=16):
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEAVY
    run.font.size = Pt(size)
    return p


def _add_matchup_table(doc, left_header, right_header, rows, metric_header=""):
    """3-col Table Grid: left_header | metric_header | right_header.

    `rows` entries are either `{"section": label}` (a subheader divider) or
    `{"label", "kind", "off_value", "off_rank", "def_value", "def_rank"}`.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    center = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell(hdr[0], left_header, HEAVY)
    _set_cell(hdr[1], metric_header, HEAVY)
    _set_cell(hdr[2], right_header, HEAVY)
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        if "section" in row:
            _set_cell(cells[0], "", HEAVY)
            _set_cell(cells[1], row["section"], HEAVY, align=center)
            _set_cell(cells[2], "", HEAVY)
            continue
        left_text = _fmt_cell(row["off_value"], row["off_rank"], row["kind"])
        right_text = _fmt_cell(row["def_value"], row["def_rank"], row["kind"])
        _set_cell(cells[0], left_text, MEDIUM, bold=True, align=center)
        _set_cell(cells[1], row["label"], HEAVY, align=center)
        _set_cell(cells[2], right_text, MEDIUM, bold=True, align=center)
    return table


def _add_schedule_table(doc, team_recap, opponent_recap):
    def team_rows(recap):
        rows = []
        for game in recap["games"]:
            prefix = "vs" if game["home_away"] == "home" else "@"
            opp_text = f"{prefix} {game['opponent']}"
            result = game["result"] or ""
            if "," in result:
                wl, score = result.split(",", 1)
                res_text = f"{score} {wl}"
            else:
                res_text = result
            rows.append((opp_text, res_text))
        return rows

    team_rows_list = team_rows(team_recap)
    opp_rows_list = team_rows(opponent_recap)
    n = max(len(team_rows_list), len(opp_rows_list))

    def header_text(recap):
        s = recap["standings"] or {}
        wins = s.get("wins", "?")
        losses = s.get("losses", "?")
        return f"{recap['team']}\nOverall: {wins}-{losses}\nConference:\n"

    table = doc.add_table(rows=1 + n, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell(hdr[0], header_text(team_recap), HEAVY, bold=True)
    _set_cell(hdr[3], header_text(opponent_recap), HEAVY, bold=True)
    for i in range(n):
        cells = table.rows[i + 1].cells
        if i < len(team_rows_list):
            _set_cell(cells[0], team_rows_list[i][0], HEAVY, bold=True)
            _set_cell(cells[1], team_rows_list[i][1], MEDIUM, bold=True)
        if i < len(opp_rows_list):
            _set_cell(cells[3], opp_rows_list[i][0], HEAVY, bold=True)
            _set_cell(cells[4], opp_rows_list[i][1], MEDIUM, bold=True)
    return table


def _add_chart_image(doc, path, caption):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc.add_picture(path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def build_weekly_report(
    conn,
    season: str,
    team: str,
    opponent: str,
    week: int,
    out_path: str,
    charts_dir: str,
) -> str:
    """Assembles the 4-page weekly coach report and saves it to `out_path`.

    Page 1: title + schedule/season recap (no Team Leaders section).
    Page 2: overall production matchup, both directions, no chart.
    Page 3: trimmed early/third-down table + one success-rate-by-down chart.
    Page 4: featured team's weekly trends, one multi-subplot chart.
    """
    from docx import Document

    charts_dir_path = Path(charts_dir)
    charts_dir_path.mkdir(parents=True, exist_ok=True)

    team_recap = rd.load_schedule_recap(conn, season, team)
    opponent_recap = rd.load_schedule_recap(conn, season, opponent)
    production = rd.load_production_matchup(conn, season, team, opponent)
    situation = rd.load_situation_trimmed(conn, season, team, opponent)
    trends = rd.load_weekly_trends(conn, season, team)

    down_chart_path = render_success_rate_by_down_chart(
        situation["success_rate_by_down"],
        team,
        opponent,
        str(charts_dir_path / "success_rate_by_down.png"),
    )
    trends_chart_path = render_weekly_trends_chart(
        trends,
        team,
        str(charts_dir_path / "weekly_trends.png"),
    )

    doc = Document()

    # --- Page 1: title + schedule/season recap ---
    title = doc.add_heading(f"Week {week} Preview: {team} vs {opponent}", level=1)
    for run in title.runs:
        run.font.name = HEAVY
    _add_schedule_table(doc, team_recap, opponent_recap)

    # --- Page 2: overall production matchup ---
    doc.add_page_break()
    _add_subheading(doc, "Overall Production", size=20)
    _add_subheading(doc, f"{team} Offense vs {opponent} Defense", size=14)
    _add_matchup_table(
        doc, f"{team} Offense", f"{opponent} Defense",
        production["team_offense_vs_opponent_defense"],
    )
    doc.add_paragraph()
    _add_subheading(doc, f"{opponent} Offense vs {team} Defense", size=14)
    _add_matchup_table(
        doc, f"{opponent} Offense", f"{team} Defense",
        production["opponent_offense_vs_team_defense"],
    )

    # --- Page 3: early & third downs, trimmed ---
    doc.add_page_break()
    _add_subheading(doc, "Early & Third Downs", size=20)
    _add_chart_image(doc, down_chart_path, "Success Rate by Down")
    _add_subheading(doc, f"{team} Offense vs {opponent} Defense", size=14)
    _add_matchup_table(
        doc, f"{team} Offense", f"{opponent} Defense",
        situation["team_offense_vs_opponent_defense"],
    )
    doc.add_paragraph()
    _add_subheading(doc, f"{opponent} Offense vs {team} Defense", size=14)
    _add_matchup_table(
        doc, f"{opponent} Offense", f"{team} Defense",
        situation["opponent_offense_vs_team_defense"],
    )

    # --- Page 4: weekly trends ---
    doc.add_page_break()
    _add_subheading(doc, f"{team} Weekly Trends", size=20)
    _add_chart_image(doc, trends_chart_path, f"{team} game-by-game trends, {season}")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
