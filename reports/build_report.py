"""
Populate the appendix of the Foothill 2025-26 Season Report with charts,
and fix known typos in the body.
Run from the repo root: python reports/build_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SRC  = r"reports\Foothill 2025-2026 Season Report.docx"
OUT  = r"reports\Foothill 2025-2026 Season Report (Final).docx"
CHARTS = r"outputs\2025-26\charts"

RED = RGBColor(0xA6, 0x1E, 0x2F)

CHART_SECTIONS = [
    ("Game-by-Game Trends", [
        ("01_success_rate_by_game.png",        "Chart 1 — Success Rate by Game"),
        ("02_explosive_rate_by_game.png",       "Chart 2 — Explosive Play Rate by Game"),
        ("03_stuff_rate_by_game.png",           "Chart 3 — Run Stuff Rate by Game"),
        ("04_third_down_by_game.png",           "Chart 4 — 3rd Down Conversion Rate by Game"),
        ("05_early_down_rush_by_game.png",      "Chart 5 — Early Down Rush Rate by Game"),
        ("06_passing_down_pass_by_game.png",    "Chart 6 — Passing Down Pass Rate by Game"),
        ("07_avg_start_pos_by_game.png",        "Chart 7 — Avg Drive Starting Position by Game"),
        ("08_penalty_yards_by_game.png",        "Chart 8 — Penalty Yards by Game"),
        ("09_half_success_by_game.png",         "Chart 9 — 1st vs. 2nd Half Success Rate by Game"),
    ]),
    ("Season Aggregates", [
        ("10_success_by_down.png",              "Chart 10 — Success Rate by Down"),
        ("11_conv_by_distance.png",             "Chart 11 — 3rd/4th Down Conversion % by Yards to Go"),
    ]),
    ("Top-15 Rankings — Early Down Defense", [
        ("12a_early_down_def_success_rt.png",   "Chart 12a — Success Rate Allowed"),
        ("12b_early_down_def_yards_per_play.png","Chart 12b — Yards per Play Allowed"),
        ("12c_early_down_def_run_stuff_pct.png","Chart 12c — Run Stuff %"),
        ("12d_early_down_def_sack_pct.png",     "Chart 12d — Sack %"),
        ("12e_early_down_def_explosive_pct.png","Chart 12e — Explosive % Allowed"),
    ]),
    ("Top-15 Rankings — Finishing Drives Defense", [
        ("13a_finishing_drives_def_pts_per_drive.png", "Chart 13a — Points per Drive Allowed"),
        ("13b_finishing_drives_def_scoring_rate.png",  "Chart 13b — Scoring Rate Allowed"),
        ("13c_finishing_drives_def_td_rate.png",       "Chart 13c — TD Rate Allowed"),
        ("13d_finishing_drives_def_avg_start_pos.png", "Chart 13d — Avg Opponent Start Position"),
    ]),
    ("Top-15 Rankings — 4th Down Offense", [
        ("14a_fourth_down_off_conv_pct.png",    "Chart 14a — Conversion %"),
        ("14b_fourth_down_off_rush_conv_pct.png","Chart 14b — Rush Conversion %"),
        ("14c_fourth_down_off_pass_pct.png",    "Chart 14c — % Pass Plays"),
        ("14d_fourth_down_off_explosives.png",  "Chart 14d — Explosive Plays"),
    ]),
    ("Top-15 Rankings — Turnovers", [
        ("15a_turnovers_off.png",               "Chart 15a — Fewest Early Down Turnovers (Offense)"),
        ("15b_turnovers_def.png",               "Chart 15b — Most 3rd Down Turnovers Forced (Defense)"),
    ]),
    ("Scatter Plots", [
        ("16_scatter_early_down_vs_scoring.png",   "Chart 16 — Early Down Success Rate vs. Drive Scoring Rate"),
        ("17_scatter_field_pos_vs_pts.png",         "Chart 17 — Avg Opponent Start Position vs. Points per Drive Allowed"),
        ("18_scatter_third_down_vs_scoring.png",    "Chart 18 — 3rd Down Conversion % vs. Drive Scoring Rate"),
        ("19_scatter_sack_vs_third_down.png",       "Chart 19 — Passing Down Sack % vs. 3rd Down Conversion % Allowed"),
        ("20_scatter_run_stuff_vs_early_down.png",  "Chart 20 — Run Stuff % vs. Early Down Success Rate Allowed"),
        ("21_scatter_turnovers_vs_pts.png",         "Chart 21 — 3rd Down Turnovers Forced vs. Points per Drive Allowed"),
    ]),
]

TYPO_FIXES = [
    ("6t-lowest",  "6th-lowest"),
    ("3r-most",    "3rd-most"),
    ("Owl’s", "Owls'"),
    ("owls defense", "Owls' defense"),
    ("The owls defense", "The Owls' defense"),
]


def fix_typos(doc):
    for para in doc.paragraphs:
        for old, new in TYPO_FIXES:
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def add_chart(doc, filename, caption):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        print(f"  MISSING: {filename}")
        return
    doc.add_picture(path, width=Inches(5.8))
    # caption
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(12)
    for run in cp.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def main():
    doc = Document(SRC)
    fix_typos(doc)

    # Find appendix heading and clear anything after it
    appendix_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Appendix" in p.text and p.style.name.startswith("Heading"):
            appendix_idx = i
            break

    if appendix_idx is None:
        print("ERROR: Could not find Appendix heading")
        return

    # Remove all paragraphs after appendix heading
    body = doc.element.body
    paras_after = doc.paragraphs[appendix_idx + 1:]
    for p in paras_after:
        body.remove(p._element)

    # Add intro line
    intro = doc.add_paragraph(
        "The following charts were generated from play-by-play data for all "
        "66 California Community College football teams in the 2025-26 season. "
        "Foothill is highlighted in red."
    )
    intro.paragraph_format.space_after = Pt(10)

    # Add each section
    for section_title, charts in CHART_SECTIONS:
        add_section_heading(doc, section_title)
        for filename, caption in charts:
            add_chart(doc, filename, caption)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
